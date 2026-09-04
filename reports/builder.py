"""
reports/builder.py
Motor comum dos relatórios em Word (.docx) — visual "premium" único para os
5 relatórios do app (Metas, Financeiro, Hábitos, Ficha de treino, Nutrição):
capa com faixa colorida, cards de KPI, barras de progresso nativas (sem
depender de imagem), tabelas com cabeçalho colorido e zebra, gráficos
matplotlib no padrão de cor do app, cabeçalho/rodapé com numeração de página.

Usa python-docx (documento) + matplotlib (gráficos, renderizados como PNG
e inseridos no documento).
"""

from __future__ import annotations

import io
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ══════════════════════════════════════════════════════════════════════════
# Paleta — mesma identidade visual do app (static/css/tokens.css), ajustada
# para fundo claro/impressão (tons um pouco mais escuros para contraste).
# ══════════════════════════════════════════════════════════════════════════

NAVY = RGBColor(0x0F, 0x17, 0x2A)
NAVY_SOFT = RGBColor(0x1E, 0x29, 0x3B)
PRIMARY = RGBColor(0x1E, 0x40, 0xAF)
PRIMARY_LIGHT = RGBColor(0x3B, 0x82, 0xF6)
PRIMARY_PALE = "EFF6FF"
SUCCESS = RGBColor(0x05, 0x96, 0x69)
SUCCESS_PALE = "ECFDF5"
DANGER = RGBColor(0xDC, 0x26, 0x26)
DANGER_PALE = "FEF2F2"
WARNING = RGBColor(0xB4, 0x53, 0x09)
WARNING_PALE = "FFFBEB"
INFO = RGBColor(0x6D, 0x28, 0xD9)
INFO_PALE = "F5F3FF"
TEXT = RGBColor(0x1E, 0x29, 0x3B)
TEXT_MUTED = RGBColor(0x64, 0x74, 0x8B)
TEXT_FAINT = RGBColor(0x94, 0xA3, 0xB8)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BORDER_HEX = "E2E8F0"
ROW_ALT_HEX = "F8FAFC"

FONT_BODY = "Calibri"

# Cores usadas nos gráficos matplotlib (hex, combinam com o app)
CHART = {
    "primary": "#1E40AF",
    "primary_light": "#3B82F6",
    "success": "#059669",
    "danger": "#DC2626",
    "warning": "#B45309",
    "info": "#6D28D9",
    "muted": "#94A3B8",
    "grid": "#E2E8F0",
    "text": "#334155",
}


# ══════════════════════════════════════════════════════════════════════════
# Helpers de baixo nível (oxml) — sombreamento de célula, bordas, margens
# ══════════════════════════════════════════════════════════════════════════

def _shade_cell(cell, hex_color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_margins(cell, top=60, bottom=60, left=120, right=120):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        node = OxmlElement(f"w:{tag}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        mar.append(node)
    tcPr.append(mar)


def _set_cell_border(cell, edges=("top", "bottom", "start", "end"), color=BORDER_HEX, sz=4):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in edges:
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)


def _no_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    tblPr.append(borders)


def _set_col_widths(table, widths_cm):
    table.autofit = False
    for row in table.rows:
        for i, w in enumerate(widths_cm):
            if i < len(row.cells):
                row.cells[i].width = Cm(w)
    for i, w in enumerate(widths_cm):
        if i < len(table.columns):
            table.columns[i].width = Cm(w)


def _bottom_border(paragraph, color=BORDER_HEX, sz=8, space=4):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_page_number_field(paragraph):
    run = paragraph.add_run()
    fld1 = OxmlElement("w:fldChar"); fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld1); run._r.append(instr); run._r.append(fld2)


# ══════════════════════════════════════════════════════════════════════════
# ReportBuilder
# ══════════════════════════════════════════════════════════════════════════

class ReportBuilder:
    def __init__(self, report_title: str, report_subtitle: str, username: str, icon: str = "📄"):
        self.doc = Document()
        self._setup_base_styles()
        self._setup_page()
        self._setup_header_footer(report_title)
        self._cover(report_title, report_subtitle, username, icon)

    # ─── Setup ──────────────────────────────────────────────────────────
    def _setup_base_styles(self):
        normal = self.doc.styles["Normal"]
        normal.font.name = FONT_BODY
        normal.font.size = Pt(10.5)
        normal.font.color.rgb = TEXT
        normal.paragraph_format.space_after = Pt(4)
        rpr = normal.element.get_or_add_rPr()
        rFonts = rpr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rpr.append(rFonts)
        rFonts.set(qn("w:eastAsia"), FONT_BODY)

    def _setup_page(self):
        section = self.doc.sections[0]
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.top_margin = Cm(1.6)
        section.bottom_margin = Cm(1.6)
        section.left_margin = Cm(1.9)
        section.right_margin = Cm(1.9)
        section.header_distance = Cm(1.0)
        section.footer_distance = Cm(1.0)

    def _setup_header_footer(self, report_title: str):
        section = self.doc.sections[0]

        header = section.header
        hp = header.paragraphs[0]
        hp.text = ""
        run = hp.add_run(f"BK Gestão Pessoal  ·  {report_title}")
        run.font.size = Pt(8.5)
        run.font.color.rgb = TEXT_FAINT
        run.font.name = FONT_BODY
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _bottom_border(hp, color=BORDER_HEX, sz=6, space=6)

        footer = section.footer
        fp = footer.paragraphs[0]
        fp.text = ""
        run1 = fp.add_run("BK Gestão Pessoal — Relatório de uso pessoal, gerado automaticamente. ")
        run1.font.size = Pt(8)
        run1.font.color.rgb = TEXT_FAINT
        run2 = fp.add_run("Página ")
        run2.font.size = Pt(8)
        run2.font.color.rgb = TEXT_FAINT
        _add_page_number_field(fp)
        run3 = fp.add_run(f" · {datetime.now().strftime('%d/%m/%Y')}")
        run3.font.size = Pt(8)
        run3.font.color.rgb = TEXT_FAINT

    def _cover(self, title: str, subtitle: str, username: str, icon: str):
        # Faixa navy full-width com título + subtítulo em branco.
        banner = self.doc.add_table(rows=1, cols=1)
        banner.alignment = WD_TABLE_ALIGNMENT.CENTER
        _no_borders(banner)
        cell = banner.cell(0, 0)
        _shade_cell(cell, "0F172A")
        _set_cell_margins(cell, top=340, bottom=340, left=340, right=340)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        p_kicker = cell.paragraphs[0]
        p_kicker.text = ""
        r = p_kicker.add_run("BK GESTÃO PESSOAL")
        r.font.size = Pt(9)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0x93, 0xC5, 0xFD)
        r.font.name = FONT_BODY

        p_title = cell.add_paragraph()
        p_title.paragraph_format.space_before = Pt(6)
        p_title.paragraph_format.space_after = Pt(4)
        r = p_title.add_run(f"{icon}  {title}")
        r.font.size = Pt(26)
        r.font.bold = True
        r.font.color.rgb = WHITE
        r.font.name = FONT_BODY

        p_sub = cell.add_paragraph()
        p_sub.paragraph_format.space_after = Pt(2)
        r = p_sub.add_run(subtitle)
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(0xBF, 0xDB, 0xFE)
        r.font.name = FONT_BODY

        self.doc.add_paragraph().paragraph_format.space_after = Pt(2)

        # Linha de metadados (usuário / data de geração)
        meta = self.doc.add_table(rows=1, cols=2)
        _no_borders(meta)
        meta.autofit = True
        left_cell, right_cell = meta.cell(0, 0), meta.cell(0, 1)
        _set_cell_margins(left_cell, 0, 0, 0, 0)
        _set_cell_margins(right_cell, 0, 0, 0, 0)

        lp = left_cell.paragraphs[0]
        r = lp.add_run(f"Preparado para: ")
        r.font.size = Pt(9.5); r.font.color.rgb = TEXT_MUTED
        r = lp.add_run(username or "—")
        r.font.size = Pt(9.5); r.font.bold = True; r.font.color.rgb = NAVY_SOFT

        rp = right_cell.paragraphs[0]
        rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = rp.add_run(f"Gerado em {datetime.now().strftime('%d/%m/%Y às %H:%M')}")
        r.font.size = Pt(9.5); r.font.color.rgb = TEXT_MUTED

        self.spacer(10)

    # ─── Elementos de conteúdo ──────────────────────────────────────────
    def spacer(self, pt=8):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(pt)
        p.paragraph_format.space_before = Pt(0)

    def section_title(self, text: str, icon: str = ""):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(8)
        _bottom_border(p, color="1E40AF", sz=14, space=6)
        label = f"{icon}  {text}".strip()
        r = p.add_run(label)
        r.font.size = Pt(15)
        r.font.bold = True
        r.font.color.rgb = NAVY
        r.font.name = FONT_BODY

    def subsection_title(self, text: str):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(text)
        r.font.size = Pt(12)
        r.font.bold = True
        r.font.color.rgb = PRIMARY
        r.font.name = FONT_BODY

    def paragraph(self, text: str, muted: bool = False, size: float = 10.5, bold: bool = False, italic: bool = False):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.italic = italic
        r.font.color.rgb = TEXT_MUTED if muted else TEXT
        r.font.name = FONT_BODY
        return p

    def callout(self, text: str, tone: str = "info", icon: str = "ℹ️"):
        """Bloco de destaque (aviso/dica) com fundo colorido suave."""
        palette = {
            "info": (INFO_PALE, INFO),
            "success": (SUCCESS_PALE, SUCCESS),
            "warning": (WARNING_PALE, WARNING),
            "danger": (DANGER_PALE, DANGER),
        }
        bg, fg = palette.get(tone, palette["info"])
        table = self.doc.add_table(rows=1, cols=1)
        _no_borders(table)
        cell = table.cell(0, 0)
        _shade_cell(cell, bg)
        _set_cell_margins(cell, top=110, bottom=110, left=160, right=160)
        p = cell.paragraphs[0]
        r = p.add_run(f"{icon}  {text}")
        r.font.size = Pt(9.8)
        r.font.color.rgb = fg
        r.font.name = FONT_BODY
        self.spacer(8)

    def kpi_row(self, items):
        """items: lista de (label, value, tone) — tone em info/success/danger/warning/primary."""
        tones = {
            "primary": (PRIMARY_PALE, PRIMARY),
            "success": (SUCCESS_PALE, SUCCESS),
            "danger": (DANGER_PALE, DANGER),
            "warning": (WARNING_PALE, WARNING),
            "info": (INFO_PALE, INFO),
        }
        n = len(items)
        table = self.doc.add_table(rows=1, cols=n)
        _no_borders(table)
        table.autofit = False
        usable_width = 21.0 - 1.9 - 1.9
        col_w = usable_width / n
        for i in range(n):
            table.columns[i].width = Cm(col_w)
        for i, (label, value, tone) in enumerate(items):
            bg, fg = tones.get(tone, tones["primary"])
            cell = table.cell(0, i)
            _shade_cell(cell, bg)
            _set_cell_margins(cell, top=140, bottom=140, left=140, right=140)
            cell.width = Cm(col_w)

            p_val = cell.paragraphs[0]
            r = p_val.add_run(str(value))
            r.font.size = Pt(18)
            r.font.bold = True
            r.font.color.rgb = fg
            r.font.name = FONT_BODY

            p_lbl = cell.add_paragraph()
            p_lbl.paragraph_format.space_before = Pt(2)
            r = p_lbl.add_run(label)
            r.font.size = Pt(8.5)
            r.font.color.rgb = TEXT_MUTED
            r.font.name = FONT_BODY
        self.spacer(10)

    def progress_bar(self, label: str, pct: float, right_text: str = None, tone: str = "primary"):
        """Barra de progresso nativa (sem imagem) via tabela de 2 células com
        largura proporcional ao percentual — some bem em Word/LibreOffice/impressão."""
        tones = {
            "primary": "1E40AF", "success": "059669", "danger": "DC2626",
            "warning": "B45309", "info": "6D28D9",
        }
        fill_hex = tones.get(tone, tones["primary"])
        pct = max(0.0, min(100.0, float(pct or 0)))

        p_lbl = self.doc.add_paragraph()
        p_lbl.paragraph_format.space_after = Pt(2)
        r = p_lbl.add_run(label)
        r.font.size = Pt(9.5); r.font.bold = True; r.font.color.rgb = NAVY_SOFT
        if right_text:
            r2 = p_lbl.add_run(f"   {right_text}")
            r2.font.size = Pt(9); r2.font.color.rgb = TEXT_MUTED

        usable_width = 21.0 - 1.9 - 1.9
        filled_w = max(0.05, usable_width * pct / 100.0)
        empty_w = max(0.0, usable_width - filled_w)

        has_empty = empty_w > 0.02
        table = self.doc.add_table(rows=1, cols=2 if has_empty else 1)
        _no_borders(table)
        table.autofit = False
        table.columns[0].width = Cm(filled_w)
        if has_empty:
            table.columns[1].width = Cm(empty_w)
        cell_fill = table.cell(0, 0)
        _shade_cell(cell_fill, fill_hex)
        _set_cell_margins(cell_fill, top=30, bottom=30, left=0, right=0)
        cell_fill.width = Cm(filled_w)
        cell_fill.paragraphs[0].add_run(" ")
        if has_empty:
            cell_empty = table.cell(0, 1)
            _shade_cell(cell_empty, "E2E8F0")
            _set_cell_margins(cell_empty, top=30, bottom=30, left=0, right=0)
            cell_empty.width = Cm(empty_w)
            cell_empty.paragraphs[0].add_run(" ")
        self.spacer(9)

    def table(self, headers, rows, col_widths=None, align=None, row_tones=None, zebra=True):
        """
        headers: list[str]
        rows: list[list[str]]
        col_widths: list[float] em cm (soma deve bater com a área útil ~17.2cm)
        align: list['left'|'right'|'center'] por coluna
        row_tones: dict {row_index: 'danger'|'warning'|'success'} para destacar linhas inteiras
        """
        n_cols = len(headers)
        align = align or ["left"] * n_cols
        row_tones = row_tones or {}
        tone_bg = {"danger": DANGER_PALE, "warning": WARNING_PALE, "success": SUCCESS_PALE}

        table = self.doc.add_table(rows=1 + len(rows), cols=n_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _no_borders(table)
        if col_widths:
            _set_col_widths(table, col_widths)

        # Cabeçalho
        for j, htext in enumerate(headers):
            cell = table.cell(0, j)
            _shade_cell(cell, "0F172A")
            _set_cell_margins(cell, top=70, bottom=70, left=100, right=100)
            p = cell.paragraphs[0]
            p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "right": WD_ALIGN_PARAGRAPH.RIGHT,
                           "center": WD_ALIGN_PARAGRAPH.CENTER}.get(align[j], WD_ALIGN_PARAGRAPH.LEFT)
            r = p.add_run(htext)
            r.font.size = Pt(9); r.font.bold = True; r.font.color.rgb = WHITE
            r.font.name = FONT_BODY

        # Linhas
        for i, row in enumerate(rows):
            tone = row_tones.get(i)
            bg = tone_bg.get(tone) if tone else (ROW_ALT_HEX if (zebra and i % 2 == 1) else None)
            for j, val in enumerate(row):
                cell = table.cell(1 + i, j)
                if bg:
                    _shade_cell(cell, bg)
                _set_cell_margins(cell, top=55, bottom=55, left=100, right=100)
                _set_cell_border(cell, edges=("bottom",), color=BORDER_HEX, sz=3)
                p = cell.paragraphs[0]
                p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "right": WD_ALIGN_PARAGRAPH.RIGHT,
                               "center": WD_ALIGN_PARAGRAPH.CENTER}.get(align[j], WD_ALIGN_PARAGRAPH.LEFT)
                r = p.add_run("" if val is None else str(val))
                r.font.size = Pt(9.3)
                r.font.name = FONT_BODY
                if tone == "danger":
                    r.font.color.rgb = DANGER
                elif tone == "warning":
                    r.font.color.rgb = WARNING
                elif tone == "success":
                    r.font.color.rgb = SUCCESS
                else:
                    r.font.color.rgb = TEXT
        self.spacer(10)

    def chart_image(self, fig, width_cm=17.2):
        buf = io.BytesIO()
        fig.savefig(buf, format="png", dpi=200, bbox_inches="tight", facecolor="white")
        buf.seek(0)
        import matplotlib.pyplot as plt
        plt.close(fig)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(buf, width=Cm(width_cm))
        self.spacer(8)

    def page_break(self):
        self.doc.add_page_break()

    def empty_state(self, text: str):
        self.callout(text, tone="info", icon="—")

    # ─── Saída ───────────────────────────────────────────────────────────
    def to_bytes(self) -> io.BytesIO:
        buf = io.BytesIO()
        self.doc.save(buf)
        buf.seek(0)
        return buf


def apply_chart_style():
    """Configura o matplotlib para combinar com o padrão visual dos relatórios."""
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    plt.rcParams.update({
        "font.family": "DejaVu Sans",
        "font.size": 9,
        "text.color": CHART["text"],
        "axes.edgecolor": CHART["grid"],
        "axes.labelcolor": CHART["text"],
        "axes.spines.top": False,
        "axes.spines.right": False,
        "xtick.color": CHART["text"],
        "ytick.color": CHART["text"],
        "grid.color": CHART["grid"],
        "figure.facecolor": "white",
        "axes.facecolor": "white",
    })


def fmt_brl(value) -> str:
    try:
        v = float(value or 0)
    except (TypeError, ValueError):
        v = 0.0
    sign = "-" if v < 0 else ""
    body = f"{abs(v):,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    return f"{sign}R$ {body}"


def nz(value):
    """Normaliza valores vindos de pandas.DataFrame.to_dict(): células vazias
    viram NaN (float), que é *truthy* em Python — `valor or "—"` não pega o
    caso. Aqui devolve None para NaN/None, preservando o valor real caso
    contrário, para então usar `nz(valor) or "—"` com segurança."""
    if value is None:
        return None
    if isinstance(value, float):
        try:
            import math
            if math.isnan(value):
                return None
        except (TypeError, ValueError):
            return None
    return value


def fmt_date(d) -> str:
    if d is None:
        return "—"
    if hasattr(d, "strftime"):
        return d.strftime("%d/%m/%Y")
    return str(d)
